import re
from pathlib import Path
from jinja2 import Template

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt


class ResumeRenderService:
    @staticmethod
    def _candidate_view(candidate: dict) -> dict:
        view = dict(candidate)
        linkedin_value = str(view.get("linkedin_url") or view.get("linkedin") or "").strip()
        github_value = str(view.get("github_url") or view.get("github") or "").strip()

        view["linkedin_url"] = linkedin_value if linkedin_value.lower().startswith(("http://", "https://")) else ""
        view["github_url"] = github_value if github_value.lower().startswith(("http://", "https://")) else ""
        view["linkedin_label"] = view.get("linkedin_label") or "LinkedIn"
        view["github_label"] = view.get("github_label") or "GitHub"
        return view

    def render_html(
        self,
        output_path: str,
        candidate: dict,
        selected_projects: list[dict],
        *,
        summary_text: str | None = None,
        skills: dict | None = None,
        experience: list[dict] | None = None,
        education: list[dict] | None = None,
        publications: list[dict] | None = None,
    ) -> str:
        rendered = self.render_html_string(
            candidate,
            selected_projects,
            summary_text=summary_text,
            skills=skills,
            experience=experience,
            education=education,
            publications=publications,
        )
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(rendered, encoding="utf-8")
        return output_path

    def render_html_string(
        self,
        candidate: dict,
        selected_projects: list[dict],
        *,
        summary_text: str | None = None,
        skills: dict | None = None,
        experience: list[dict] | None = None,
        education: list[dict] | None = None,
        publications: list[dict] | None = None,
    ) -> str:
        template_path = Path("templates/resume_template.html")
        template = Template(template_path.read_text(encoding="utf-8"))
        return template.render(
            candidate=self._candidate_view(candidate),
            projects=selected_projects,
            summary_text=summary_text,
            skills=skills or {},
            experience=experience or [],
            education=education or [],
            publications=publications or [],
        )

    def render_docx(
        self,
        output_path: str,
        candidate: dict,
        selected_projects: list[dict],
        *,
        summary_text: str | None = None,
        skills: dict | None = None,
        experience: list[dict] | None = None,
        education: list[dict] | None = None,
        publications: list[dict] | None = None,
    ) -> str:
        doc = Document()
        self._configure_docx(doc)

        candidate_view = self._candidate_view(candidate)
        self._add_header(doc, candidate_view)
        self._add_summary(doc, summary_text or candidate_view.get("base_summary", ""))
        self._add_section(doc, "Technical Skills")
        self._add_skills(doc, skills or {})
        self._add_section(doc, "Professional Experience")
        self._add_experience(doc, experience or [])
        if selected_projects:
            self._add_section(doc, "Key Projects")
            self._add_projects(doc, selected_projects)
        if publications:
            self._add_section(doc, "Research & Publications")
            for item in publications:
                self._add_role_line(
                    doc,
                    f"{item.get('title', '')}, {item.get('venue', '')}",
                    str(item.get("year") or ""),
                    size=7.8,
                )
                for bullet in (item.get("bullets") or []):
                    self._add_bullet(doc, bullet, size=7.8)
        self._add_section(doc, "Education")
        self._add_education(doc, education or [])

        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        return output_path

    @staticmethod
    def _configure_docx(doc: Document) -> None:
        section = doc.sections[0]
        section.top_margin = Inches(0.33)
        section.bottom_margin = Inches(0.33)
        section.left_margin = Inches(0.42)
        section.right_margin = Inches(0.42)

        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(8.1)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.0

        for style_name in ("List Bullet", "List Paragraph"):
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.size = Pt(8.0)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.line_spacing = 1.0
            style.paragraph_format.left_indent = Inches(0.16)
            style.paragraph_format.first_line_indent = Inches(-0.09)

    def _add_header(self, doc: Document, candidate: dict) -> None:
        name = doc.add_paragraph()
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.paragraph_format.space_after = Pt(0)
        run = name.add_run(self._resume_text(candidate.get("full_name", "")).upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14.5)

        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.paragraph_format.space_after = Pt(1)
        contact_parts = [
            candidate.get("location", ""),
            candidate.get("email", ""),
            candidate.get("phone", ""),
        ]
        self._add_inline(contact, " | ".join(item for item in contact_parts if item), size=7.4)
        if candidate.get("linkedin_url"):
            self._add_inline(contact, " | ", size=7.4, preserve_spacing=True)
            self._add_hyperlink(contact, "LinkedIn", candidate["linkedin_url"], size=7.4)
        if candidate.get("github_url"):
            self._add_inline(contact, " | ", size=7.4, preserve_spacing=True)
            self._add_hyperlink(contact, "GitHub", candidate["github_url"], size=7.4)

    def _add_summary(self, doc: Document, value: str) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        self._add_inline(paragraph, self._resume_text(value), size=7.8)

    def _add_section(self, doc: Document, title: str) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(title.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.8)
        self._add_bottom_border(paragraph)

    def _add_skills(self, doc: Document, skills: dict) -> None:
        for group, values in skills.items():
            if not values:
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            label = group.replace("_", " ").replace(" and ", " & ").title()
            label_run = paragraph.add_run(f"{self._resume_text(label)}: ")
            label_run.bold = True
            label_run.font.name = "Arial"
            label_run.font.size = Pt(7.7)
            self._add_inline(paragraph, ", ".join(self._resume_text(item) for item in values), size=7.7)

    def _add_experience(self, doc: Document, experience: list[dict]) -> None:
        for item in experience:
            self._add_role_line(
                doc,
                f"{item.get('title', '')}, {item.get('company', '')}",
                item.get("dates", ""),
            )
            if item.get("location"):
                location = doc.add_paragraph()
                self._add_inline(location, item.get("location", ""), size=7.3, italic=True)
            for bullet in (item.get("bullets") or []):
                self._add_bullet(doc, bullet, size=7.8)

    def _add_projects(self, doc: Document, projects: list[dict]) -> None:
        for project in projects[:3]:
            tech = ", ".join(self._resume_text(item) for item in project.get("tech_stack", []))
            title = project.get("name", "")
            self._add_role_line(doc, f"{title}{f' | {tech}' if tech else ''}", "")
            for bullet in (project.get("bullets") or []):
                self._add_bullet(doc, bullet, size=7.8)

    def _add_education(self, doc: Document, education: list[dict]) -> None:
        for item in education:
            degree = item.get("degree", "")
            institution = item.get("institution", "")
            gpa = f" | GPA: {item.get('gpa')}" if item.get("gpa") else ""
            date = item.get("dates") or item.get("expected_graduation") or ""
            self._add_role_line(doc, f"{degree}, {institution}{gpa}", date, size=7.8)

    def _add_role_line(self, doc: Document, left: str, right: str, *, size: float = 7.9) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.55), WD_TAB_ALIGNMENT.RIGHT)
        left_run = paragraph.add_run(self._resume_text(left))
        left_run.bold = True
        left_run.font.name = "Arial"
        left_run.font.size = Pt(size)
        if right:
            right_run = paragraph.add_run(f"\t{self._resume_text(right)}")
            right_run.font.name = "Arial"
            right_run.font.size = Pt(size)

    def _add_bullet(self, doc: Document, value: str, *, size: float) -> None:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0)
        self._add_inline(paragraph, self._resume_text(value), size=size)

    @staticmethod
    def _add_inline(
        paragraph,
        value: object,
        *,
        size: float,
        bold: bool = False,
        italic: bool = False,
        preserve_spacing: bool = False,
    ):
        text = ResumeRenderService._resume_text(value, preserve_spacing=preserve_spacing)
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if preserve_spacing and text != text.strip():
            for text_node in run._r.iter(qn("w:t")):
                text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        return run

    @staticmethod
    def _add_hyperlink(paragraph, text: str, url: str, *, size: float) -> None:
        relationship_id = paragraph.part.relate_to(
            url,
            RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        font = OxmlElement("w:rFonts")
        font.set(qn("w:ascii"), "Arial")
        font.set(qn("w:hAnsi"), "Arial")
        run_properties.append(font)

        size_node = OxmlElement("w:sz")
        size_node.set(qn("w:val"), str(int(size * 2)))
        run_properties.append(size_node)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(underline)

        run.append(run_properties)
        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    @staticmethod
    def _add_bottom_border(paragraph) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "555555")
        p_bdr.append(bottom)

    @staticmethod
    def _resume_text(value: object, *, preserve_spacing: bool = False) -> str:
        text = "" if value is None else str(value)
        replacements = {
            "\u2014": ", ",
            "\u2013": " to ",
            "\u2212": " to ",
            "\u00a0": " ",
            "<": "",
            ">": "",
            "~": "",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        text = text.replace(" - ", ", ")
        return re.sub(r"\s+", " ", text) if preserve_spacing else " ".join(text.split())
